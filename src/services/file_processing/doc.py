import os
from docx import Document


class DocProcessor:
    def parse(self, file_path, page_number=0):
        self.file_path = file_path
        self.file_extension = os.path.splitext(file_path)[1].lower()
        self.page_number = page_number
        self.file_name = os.path.basename(file_path)

        if self.file_extension == ".docx":
            return self._parse_docx()
        else:
            raise ValueError("Unsupported file format")

    def _parse_docx(self):
        document = Document(self.file_path)
        text = []
        tables_data = []

        # Python-docx does not support page numbers, so we'll assume the whole document as one page
        for para in document.paragraphs:
            text.append(para.text)

        for i, table in enumerate(document.tables):
            table_data = [[cell.text for cell in row.cells] for row in table.rows]
            tables_data.append(table_data)

        text_data = {
            "raw_text": " ".join(text),
            "file_path": self.file_path,
            "doc_type": "text",
            "filename": self.file_name,
            "page_number": self.page_number,
        }

        tables_json_data = [
            {
                "raw_text": "",
                "file_path": self.file_path,
                "doc_type": "docx_table",
                "raw_table": table_data,
                "filename": self.file_name,
                "page_number": self.page_number,
            }
            for table_data in tables_data
        ]

        return text_data, tables_json_data
