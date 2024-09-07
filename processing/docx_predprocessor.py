import os
import json
from docx import Document
import pythoncom
from win32com.client import Dispatch

class DocParser:
    def __init__(self):
        pass
    
    def parse(self, file_path, page_number):
        self.file_path = file_path
        self.file_extension = os.path.splitext(file_path)[1].lower()
        self.page_number = page_number
        self.file_name = os.path.basename(file_path)

        if self.file_extension == '.docx':
            return self._parse_docx()
        elif self.file_extension == '.doc':
            return self._parse_doc()
        else:
            raise ValueError("Unsupported file format")
    
    def _parse_docx(self):
        document = Document(self.file_path)
        text = []
        tables_data = []

        # Мы предполагаем, что python-docx не поддерживает извлечение по страницам,
        # поэтому предполагаем, что весь документ - это одна страница.

        for para in document.paragraphs:
            text.append(para.text)

        for i, table in enumerate(document.tables):
            table_data = [[cell.text for cell in row.cells] for row in table.rows]
            tables_data.append(table_data)

        text_data = {
            "raw_text": ' '.join(text),
            "file_path": self.file_path,
            "doc_type": 'text',
            "filename": self.file_name,
            "page_number": self.page_number
        }

        tables_json_data = [
            {
                "raw_text": '',
                "file_path": self.file_path,
                "doc_type": 'docx_table',
                "raw_table": table_data,
                "filename": self.file_name,
                "page_number": self.page_number
            }
            for table_data in tables_data
        ]

        return text_data, tables_json_data

    def _parse_doc(self):
        pythoncom.CoInitialize()
        word = Dispatch("Word.Application")
        doc = word.Documents.Open(self.file_path)
        text = []
        tables_data = []

        # Получаем текст и таблицы с конкретной страницы
        page_text, page_tables = self._get_text_and_tables_from_page(doc, self.page_number)
        text.extend(page_text)
        tables_data.extend(page_tables)
        
        doc.Close()
        word.Quit()
        pythoncom.CoUninitialize()
        
        text_data = {
            "raw_text": ' '.join(text),
            "file_path": self.file_path,
            "doc_type": 'text',
            "filename": self.file_name,
            "page_number": self.page_number
        }

        tables_json_data = [
            {
                "raw_text": '',
                "file_path": self.file_path,
                "doc_type": 'doc_table',
                "raw_table": table_data,
                "filename": self.file_name,
                "page_number": self.page_number
            }
            for table_data in tables_data
        ]

        return text_data, tables_json_data

    def _get_text_and_tables_from_page(self, doc, page_number):
        page_text = []
        page_tables = []

        # Получаем текст с конкретной страницы
        for para in doc.Paragraphs:
            if para.Range.Information(3) == page_number:
                page_text.append(para.Range.Text.strip())
        
        # Получаем таблицы с конкретной страницы
        previous_headers = None
        for table in doc.Tables:
            if table.Range.Information(3) == page_number:
                table_data = [[cell.Range.Text.strip() for cell in row.Cells] for row in table.Rows]
                
                if previous_headers and self._are_headers_equal(previous_headers, table_data[0]):
                    # Если заголовки совпадают, это продолжение предыдущей таблицы
                    page_tables[-1].extend(table_data[1:])  # Добавляем все, кроме заголовков
                else:
                    page_tables.append(table_data)
                
                previous_headers = table_data[0]
        
        return page_text, page_tables

    def _are_headers_equal(self, headers1, headers2):
        if len(headers1) != len(headers2):
            return False
        return all(cell1 == cell2 for cell1, cell2 in zip(headers1, headers2))

